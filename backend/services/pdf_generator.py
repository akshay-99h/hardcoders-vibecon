"""
PDF Generation Service for RakshaAI Documents
Enhanced professional formatting with official government-style appearance
"""
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib import colors
from io import BytesIO
from datetime import datetime
import re
import uuid


class PDFGeneratorService:
    """Service to generate professional PDF documents with official appearance"""

    _MARKDOWN_BOLD_PATTERN = re.compile(
        r"\*\*(?=\S)(.+?)(?<=\S)\*\*|\*(?=\S)(.+?)(?<=\S)\*"
    )

    @staticmethod
    def _format_markdown_for_pdf(text: str) -> str:
        """Escape XML entities and convert markdown asterisks to reportlab bold tags."""
        escaped = (
            text.replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
        )

        def _replace(match):
            bold_text = match.group(1) or match.group(2) or ""
            return f"<b>{bold_text}</b>"

        return PDFGeneratorService._MARKDOWN_BOLD_PATTERN.sub(_replace, escaped)

    @staticmethod
    def _split_markdown_bold_segments(text: str):
        """Split a line into plain/bold segments from markdown asterisk syntax."""
        segments = []
        last_index = 0

        for match in PDFGeneratorService._MARKDOWN_BOLD_PATTERN.finditer(text):
            start, end = match.span()
            if start > last_index:
                segments.append((text[last_index:start], False))

            bold_text = match.group(1) or match.group(2) or ""
            if bold_text:
                segments.append((bold_text, True))
            last_index = end

        if last_index < len(text):
            segments.append((text[last_index:], False))

        return segments if segments else [(text, False)]

    @staticmethod
    def _add_markdown_line_to_docx(paragraph, text: str, *, force_bold: bool = False):
        """Append a line to a DOCX paragraph and honor markdown asterisk bold markers."""
        for segment, is_bold in PDFGeneratorService._split_markdown_bold_segments(text):
            if not segment:
                continue
            run = paragraph.add_run(segment)
            if force_bold or is_bold:
                run.bold = True
    
    @staticmethod
    def _add_header_footer(canvas_obj, doc):
        """Add official-looking header and footer to each page"""
        canvas_obj.saveState()
        
        width, height = A4
        
        # Header - Official document marker
        canvas_obj.setStrokeColor(colors.HexColor('#003366'))
        canvas_obj.setLineWidth(2)
        canvas_obj.line(0.75*inch, height - 0.5*inch, width - 0.75*inch, height - 0.5*inch)
        
        # Reference number (top right corner)
        canvas_obj.setFont('Times-Roman', 9)
        canvas_obj.setFillColor(colors.HexColor('#666666'))
        ref_num = f"Ref: RAI/{datetime.now().strftime('%Y%m')}/{uuid.uuid4().hex[:8].upper()}"
        canvas_obj.drawRightString(width - 0.75*inch, height - 0.35*inch, ref_num)
        
        # Footer - Page number
        canvas_obj.setFont('Times-Roman', 9)
        canvas_obj.setFillColor(colors.HexColor('#666666'))
        page_text = f"Page {doc.page}"
        canvas_obj.drawCentredString(width/2, 0.5*inch, page_text)
        
        # Footer line
        canvas_obj.setStrokeColor(colors.HexColor('#cccccc'))
        canvas_obj.setLineWidth(0.5)
        canvas_obj.line(0.75*inch, 0.65*inch, width - 0.75*inch, 0.65*inch)
        
        canvas_obj.restoreState()
    
    @staticmethod
    def generate_document_pdf(
        document_type: str,
        document_content: str,
        user_name: str = "Citizen"
    ) -> BytesIO:
        """
        Generate a professional PDF from document content with official government-style formatting
        
        Args:
            document_type: Type of document (RTI, Complaint, etc.)
            document_content: Full text content of the document
            user_name: Name of the user generating the document
            
        Returns:
            BytesIO buffer containing the PDF
        """
        buffer = BytesIO()
        
        # Create PDF with A4 page size and proper margins
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=1*inch,
            leftMargin=1*inch,
            topMargin=1*inch,
            bottomMargin=1*inch,
        )
        
        # Add header/footer to each page
        doc.build = lambda story: SimpleDocTemplate.build(
            doc, story, 
            onFirstPage=PDFGeneratorService._add_header_footer,
            onLaterPages=PDFGeneratorService._add_header_footer
        )
        
        # Container for PDF elements
        elements = []
        
        # Define professional styles using Times-Roman for more formal appearance
        styles = getSampleStyleSheet()
        
        # Document title style
        title_style = ParagraphStyle(
            'DocumentTitle',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#000000'),
            spaceAfter=0.3*inch,
            alignment=TA_CENTER,
            fontName='Times-Bold',
            leading=18
        )
        
        # Address style (To, From) - Official format
        address_style = ParagraphStyle(
            'Address',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#000000'),
            spaceAfter=6,
            alignment=TA_LEFT,
            fontName='Times-Roman',
            leading=15
        )
        
        # Subject line style - Bold and underlined for emphasis
        subject_style = ParagraphStyle(
            'Subject',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#000000'),
            spaceAfter=14,
            spaceBefore=14,
            alignment=TA_LEFT,
            fontName='Times-Bold',
            leading=15
        )
        
        # Body paragraph style - Justified for formal appearance
        body_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#000000'),
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            fontName='Times-Roman',
            leading=18,
            firstLineIndent=0
        )
        
        # Salutation style
        salutation_style = ParagraphStyle(
            'Salutation',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#000000'),
            spaceAfter=14,
            alignment=TA_LEFT,
            fontName='Times-Roman',
            leading=15
        )
        
        # Signature block style
        signature_style = ParagraphStyle(
            'Signature',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#000000'),
            spaceAfter=7,
            alignment=TA_LEFT,
            fontName='Times-Roman',
            leading=15
        )
        
        # Date style (top right) - Official format
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#000000'),
            alignment=TA_RIGHT,
            fontName='Times-Roman'
        )
        
        # Add date at top right
        date_text = f"Date: {datetime.now().strftime('%d %B %Y')}"
        elements.append(Paragraph(date_text, date_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Process document content
        lines = document_content.split('\n')
        
        in_address_block = False
        in_subject = False
        in_body = False
        in_signature = False
        
        for line in lines:
            raw_line = line.strip()

            if not raw_line:
                # Empty line - add small spacer
                elements.append(Spacer(1, 0.1*inch))
                continue

            # Detect different sections
            line_lower = raw_line.lower()
            formatted_line = PDFGeneratorService._format_markdown_for_pdf(raw_line)
            
            # To/From address blocks
            if line_lower.startswith('to,') or line_lower.startswith('to:'):
                in_address_block = True
                elements.append(Paragraph(formatted_line, address_style))
            
            # Subject line
            elif line_lower.startswith('subject:'):
                in_subject = True
                in_address_block = False
                elements.append(Paragraph(formatted_line, subject_style))
            
            # Salutations
            elif any(sal in line_lower for sal in ['respected sir', 'dear sir', 'sir/madam', 'dear madam']):
                in_body = True
                in_subject = False
                elements.append(Paragraph(formatted_line, salutation_style))
            
            # Closing (Yours faithfully, Thanking you, etc.)
            elif any(closing in line_lower for closing in ['yours faithfully', 'yours sincerely', 'thanking you', 'regards']):
                in_body = False
                in_signature = True
                elements.append(Spacer(1, 0.15*inch))
                elements.append(Paragraph(formatted_line, signature_style))
            
            # Enclosures
            elif line_lower.startswith('enclosure'):
                in_signature = True
                elements.append(Spacer(1, 0.15*inch))
                elements.append(Paragraph(f"<b>{formatted_line}</b>", signature_style))
            
            # Address block lines
            elif in_address_block:
                elements.append(Paragraph(formatted_line, address_style))
            
            # Signature block (name, address, phone)
            elif in_signature:
                elements.append(Paragraph(formatted_line, signature_style))
            
            # Body paragraphs
            else:
                # Check if it's a heading/subheading
                if (len(raw_line) < 60 and raw_line.isupper()) or raw_line.endswith(':'):
                    heading_style = ParagraphStyle(
                        'Heading',
                        parent=styles['Heading3'],
                        fontSize=12,
                        textColor=colors.HexColor('#000000'),
                        spaceAfter=10,
                        spaceBefore=14,
                        fontName='Times-Bold',
                        leading=15
                    )
                    elements.append(Paragraph(formatted_line, heading_style))
                else:
                    # Regular body text
                    elements.append(Paragraph(formatted_line, body_style))
        
        # Add signature space
        elements.append(Spacer(1, 0.5*inch))
        
        # Add a box for signature and stamp
        signature_box_style = ParagraphStyle(
            'SignatureBox',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#666666'),
            alignment=TA_RIGHT,
            fontName='Times-Roman'
        )
        
        # Signature placeholder
        signature_table = Table(
            [
                [''],
                ['(Signature)'],
                [''],
                ['Place for Stamp/Seal']
            ],
            colWidths=[2.5*inch],
            style=TableStyle([
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#999999')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#666666')),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ])
        )
        
        # Right-align the signature box
        elements.append(Table(
            [[signature_table]],
            colWidths=[6.5*inch],
            style=TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ])
        ))
        
        # Add professional footer section
        elements.append(Spacer(1, 0.4*inch))
        
        # Separator line - more prominent
        elements.append(Table(
            [['']], 
            colWidths=[6.5*inch],
            style=TableStyle([
                ('LINEABOVE', (0, 0), (-1, 0), 1, colors.HexColor('#003366')),
            ])
        ))
        
        elements.append(Spacer(1, 0.15*inch))
        
        # Instructions box - more formal
        instruction_style = ParagraphStyle(
            'Instruction',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#333333'),
            alignment=TA_LEFT,
            leftIndent=10,
            spaceAfter=5,
            fontName='Times-Roman'
        )
        
        instructions = [
            "<b>FILING INSTRUCTIONS:</b>",
            "1. Print this document on standard A4 white paper (70-80 GSM)",
            "2. Sign above the printed name where indicated",
            "3. Affix stamp/seal if applicable",
            "4. Attach all supporting documents as listed in 'Enclosures'",
            "5. Retain one photocopy for personal records with postal receipt",
            "6. Send via Registered Post/Speed Post to the address mentioned above",
            "7. Note the postal tracking number for future reference"
        ]
        
        for instruction in instructions:
            elements.append(Paragraph(instruction, instruction_style))
        
        # Official disclaimer
        elements.append(Spacer(1, 0.25*inch))
        
        # Disclaimer box with border
        disclaimer_style = ParagraphStyle(
            'Disclaimer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#555555'),
            alignment=TA_JUSTIFY,
            fontName='Times-Italic',
            leading=11,
            leftIndent=5,
            rightIndent=5
        )
        
        disclaimer_text = (
            "<b>DISCLAIMER:</b> This document has been generated by RakshaAI, an AI-powered assistant. "
            "This is a draft template and should be carefully reviewed before submission. "
            "The user is responsible for verifying all information, dates, addresses, and legal compliance. "
            "For complex legal matters or professional advice, please consult a licensed advocate. "
            "<b>Free Legal Aid:</b> NALSA Helpline 15100 | www.nalsa.gov.in | "
            "<b>RTI Portal:</b> rtionline.gov.in"
        )
        
        disclaimer_table = Table(
            [[Paragraph(disclaimer_text, disclaimer_style)]],
            colWidths=[6.5*inch],
            style=TableStyle([
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#999999')),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f5f5f5')),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ])
        )
        elements.append(disclaimer_table)
        
        # Build PDF
        doc.build(elements)
        
        # Reset buffer position
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_document_docx(
        document_type: str,
        document_content: str,
        user_name: str = "Citizen"
    ) -> BytesIO:
        """
        Generate a DOCX document from text content, preserving markdown asterisk bold markers.
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
        except ModuleNotFoundError as exc:
            raise RuntimeError(
                "DOCX generation requires the 'python-docx' package. Install backend requirements first."
            ) from exc

        doc = Document()

        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)

        title = doc.add_paragraph(document_type.replace('_', ' ').strip() or "Document")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(14)

        if user_name:
            generated_for = doc.add_paragraph(f"Prepared for: {user_name}")
            generated_for.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        date_paragraph = doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph("")

        lines = document_content.split('\n')
        in_address_block = False
        in_signature = False

        for line in lines:
            raw_line = line.strip()
            if not raw_line:
                doc.add_paragraph("")
                continue

            line_lower = raw_line.lower()
            paragraph = doc.add_paragraph()

            if line_lower.startswith('to,') or line_lower.startswith('to:'):
                in_address_block = True
                in_signature = False
                PDFGeneratorService._add_markdown_line_to_docx(paragraph, raw_line)
                continue

            if line_lower.startswith('subject:'):
                in_address_block = False
                PDFGeneratorService._add_markdown_line_to_docx(paragraph, raw_line, force_bold=True)
                continue

            if any(sal in line_lower for sal in ['respected sir', 'dear sir', 'sir/madam', 'dear madam']):
                in_signature = False
                PDFGeneratorService._add_markdown_line_to_docx(paragraph, raw_line)
                continue

            if any(closing in line_lower for closing in ['yours faithfully', 'yours sincerely', 'thanking you', 'regards']):
                in_signature = True
                PDFGeneratorService._add_markdown_line_to_docx(paragraph, raw_line)
                continue

            if line_lower.startswith('enclosure'):
                in_signature = True
                PDFGeneratorService._add_markdown_line_to_docx(paragraph, raw_line, force_bold=True)
                continue

            if in_address_block or in_signature:
                PDFGeneratorService._add_markdown_line_to_docx(paragraph, raw_line)
                continue

            if (len(raw_line) < 60 and raw_line.isupper()) or raw_line.endswith(':'):
                PDFGeneratorService._add_markdown_line_to_docx(paragraph, raw_line, force_bold=True)
            else:
                PDFGeneratorService._add_markdown_line_to_docx(paragraph, raw_line)

        doc.add_paragraph("")
        disclaimer = doc.add_paragraph(
            "DISCLAIMER: This document has been generated by RakshaAI. "
            "Please review all details before submission."
        )
        if disclaimer.runs:
            disclaimer.runs[0].italic = True

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
